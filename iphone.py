import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import tempfile
import datetime
from PIL import Image
import cv2

# --- Smart Helper to find Logo File in Repository ---
def get_logo_path():
    for file in os.listdir('.'):
        if file.lower().startswith('logo') and file.lower().endswith(('.png', '.jpg', '.jpeg')):
            return file
    return None

# --- Function to extract thumbnail frame from Video ---
def extract_video_frame(video_file):
    try:
        # Save video temporarily to read frame
        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as tmp:
            tmp.write(video_file.getvalue())
            tmp_path = tmp.name

        cap = cv2.VideoCapture(tmp_path)
        ret, frame = cap.read() # Capture 1st frame
        cap.release()
        os.remove(tmp_path)

        if ret:
            frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            pil_img = Image.fromarray(frame_rgb)
            img_byte_arr = io.BytesIO()
            pil_img.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            return img_byte_arr
    except Exception as e:
        st.error(f"Video frame extract error: {e}")
    return None

# --- Word Document Generator Function ---
def create_inspection_report(images_data, videos_data, all_sections, details):
    doc = Document()
    
    # 1. Add Logo in Document HEADER (Top Left Side)
    logo_file = get_logo_path()
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if logo_file:
        try:
            header_para.add_run().add_picture(logo_file, width=Inches(1.8))
        except Exception as e:
            header_para.text = "LIFELONG QUALITY INSPECTION"
    else:
        header_para.text = "LIFELONG QUALITY INSPECTION"

    # 2. Document Title
    doc.add_heading('Quality Inspection Report', 0)
    
    # 3. Add Inspection Details Table
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = 'Table Grid'
    
    detail_items = [
        ("Inspection Date", str(details['date'])),
        ("Vendor Name", details['vendor']),
        ("SKU ID", details['sku']),
        ("Produced Qty", str(details['produced'])),
        ("Inspected Qty", str(details['inspected'])),
        ("Lot Status", details['status'])
    ]
    
    for i, (key, val) in enumerate(detail_items):
        row_cells = meta_table.rows[i].cells
        row_cells[0].text = key
        row_cells[1].text = val if val else "N/A"
        row_cells[0].paragraphs[0].runs[0].bold = True
        
    doc.add_paragraph() # Spacing
    
    # 4. Add Evidence Sections (Images + Video Snapshots inside Doc Grid)
    for sec_name in all_sections:
        images = images_data.get(sec_name, [])
        videos = videos_data.get(sec_name, [])
        
        # Combine photos and video frames into a single media list
        combined_media = []
        for img in images:
            combined_media.append(("📷 Photo", img))
            
        for vid in videos:
            frame_bytes = extract_video_frame(vid)
            if frame_bytes:
                combined_media.append((f"🎥 Video Frame ({vid.name})", frame_bytes))
        
        num_cols = 2 
        num_rows = max(1, (len(combined_media) + num_cols - 1) // num_cols) if combined_media else 1
        
        table = doc.add_table(rows=num_rows + 1, cols=num_cols)
        table.style = 'Table Grid' 
        
        # Section Header Row
        header_cell = table.cell(0, 0)
        header_cell.merge(table.cell(0, num_cols - 1))
        header_paragraph = header_cell.paragraphs[0]
        run = header_paragraph.add_run(sec_name)
        run.bold = True 
        
        if combined_media:
            for i, (label, item_bytes) in enumerate(combined_media):
                row_idx = (i // num_cols) + 1
                col_idx = i % num_cols
                cell = table.cell(row_idx, col_idx)
                paragraph = cell.paragraphs[0]
                
                run = paragraph.add_run()
                if hasattr(item_bytes, 'getvalue'):
                    run.add_picture(io.BytesIO(item_bytes.getvalue()), width=Inches(2.5))
                else:
                    run.add_picture(io.BytesIO(item_bytes.read()), width=Inches(2.5))
                
                # Label below the media item
                p_label = cell.add_paragraph(label)
                p_label.runs[0].font.size = Inches(0.11)
        else:
            empty_cell = table.cell(1, 0)
            empty_cell.merge(table.cell(1, num_cols - 1))
            empty_cell.text = "No images or videos provided"
        
        doc.add_paragraph()
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- Streamlit UI Design ---
st.title("📱 Lifelong QA Report Maker")

if 'selected_images' not in st.session_state:
    st.session_state.selected_images = {}
if 'selected_videos' not in st.session_state:
    st.session_state.selected_videos = {}
if 'sections_list' not in st.session_state:
    st.session_state.sections_list = ['MRP', 'Barcode (EAN Code)', 'Outer Box', 'Inner Box', 'Drop Test', 'Functional Testing', 'Visual Inspection']

# Step 1: Inspection Details Form
st.subheader("1. Inspection Details")
col1, col2 = st.columns(2)
with col1:
    insp_date = st.date_input("Inspection Date", datetime.date.today())
    vendor_name = st.text_input("Vendor Name")
    sku_id = st.text_input("SKU ID")
with col2:
    produced_qty = st.number_input("Produced Qty", min_value=0)
    inspected_qty = st.number_input("Inspected Qty", min_value=0)
    lot_status = st.selectbox("Lot Status", ["Pending", "Accepted", "Rejected"])

st.divider()

# Step 2: Manage Sections
st.subheader("2. Select or Add Section")
section_choice = st.selectbox("Choose section:", st.session_state.sections_list)
with st.expander("Or Add Custom Section"):
    new_section = st.text_input("Section Name")
    if st.button("Add Section"):
        if new_section and new_section not in st.session_state.sections_list:
            st.session_state.sections_list.append(new_section)
            st.success(f"Added '{new_section}'!")
            st.rerun()

# Step 3: Upload Photos & Videos
st.subheader("3. Upload Evidence")
uploaded_files = st.file_uploader(f"Upload photos & videos for {section_choice}", type=['jpg', 'jpeg', 'png', 'mp4', 'mov'], accept_multiple_files=True)
if st.button("Save Files to Section", type="primary"):
    if uploaded_files:
        if section_choice not in st.session_state.selected_images:
            st.session_state.selected_images[section_choice] = []
        if section_choice not in st.session_state.selected_videos:
            st.session_state.selected_videos[section_choice] = []
            
        for file in uploaded_files:
            if file.name.lower().endswith(('.mp4', '.mov')):
                st.session_state.selected_videos[section_choice].append(file)
            else:
                st.session_state.selected_images[section_choice].append(file)
                
        st.success(f"Added {len(uploaded_files)} file(s) to {section_choice}!")
    else:
        st.error("Please select files first.")

# Step 4: Live Preview
if st.session_state.selected_images or st.session_state.selected_videos:
    st.subheader("4. Evidence Preview")
    for sec in st.session_state.sections_list:
        imgs = st.session_state.selected_images.get(sec, [])
        vids = st.session_state.selected_videos.get(sec, [])
        
        if imgs or vids:
            st.write(f"**{sec}** ({len(imgs)} photos, {len(vids)} videos)")
            if imgs:
                cols = st.columns(2)
                for i, img in enumerate(imgs):
                    with cols[i % 2]:
                        st.image(img, use_container_width=True)
            if vids:
                for vid in vids:
                    st.info(f"🎥 Video: {vid.name}")

# Step 5: Generate Report
st.divider()
st.subheader("5. Finalize Report")
report_name = st.text_input("Report Name", f"{sku_id}_Inspection_Report" if sku_id else "Quality_Inspection_Report")

if st.button("Generate Word Report", type="primary"):
    if not st.session_state.selected_images and not st.session_state.selected_videos:
        st.warning("Please add at least one image or video before generating the report.")
    else:
        with st.spinner("Creating Word Report with images & video frames..."):
            
            report_details = {
                'date': insp_date,
                'vendor': vendor_name,
                'sku': sku_id,
                'produced': produced_qty,
                'inspected': inspected_qty,
                'status': lot_status
            }
            
            docx_data = create_inspection_report(st.session_state.selected_images, st.session_state.selected_videos, st.session_state.sections_list, report_details)
            
            st.download_button(
                label="📥 Download .docx Word Report",
                data=docx_data,
                file_name=f"{report_name}.docx",
                mime="application/octet-stream"
            )

if st.button("Reset / Clear All Data"):
    st.session_state.selected_images = {}
    st.session_state.selected_videos = {}
    st.rerun()
